from io import BytesIO
import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from utils.questions import general_questions, shipping_logistics_questions, bid_rfi_rfp_questions

def get_docx_file(questions, answers_list, files_pages_list, processed_file_type, selected_file_type):
    try:
        if processed_file_type == "General" and questions != general_questions:
            st.warning(f"The answers below are for the **{processed_file_type}** category. Please click on **Generate Answers** button to generate answers for the **{selected_file_type}** category!")

        elif processed_file_type == "Shipping & Logistics" and questions != shipping_logistics_questions:
            st.warning(f"The answers below are for the **{processed_file_type}** category. Please click on **Generate Answers** button to generate answers for the **{selected_file_type}** category!")

        elif processed_file_type == "RFI/RFP (Bid)" and questions != bid_rfi_rfp_questions:
            st.warning(f"The answers below are for the **{processed_file_type}** category. Please click on **Generate Answers** button to generate answers for the **{selected_file_type}** category!")
        
        else:
            # Create a new Document
            doc = Document()
            doc.add_heading("PDF QUESTION ANSWERING REPORT", level=1)
            #doc.add_paragraph("=" * 100)
            doc.add_paragraph("")  # Add some spacing
            
            # Add content
            iter = 1
            for question in questions:
                # Add question heading
                question_heading = doc.add_heading(f"QUESTION {iter}: {question.upper()}", level=2)
                question_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Add answer
                answer_paragraph = doc.add_paragraph(f"ANSWER: {answers_list[iter-1]}")
                answer_paragraph.style.font.size = Pt(12)
                
                # Add files and pages referenced
                doc.add_paragraph("FILES AND PAGES REFERENCED:") #, style="Intense Quote")
                for file_name, page_num in files_pages_list[iter-1]:
                    doc.add_paragraph(f"File: {file_name}, Page: {page_num}")
                
                # Add separator
                #doc.add_paragraph("=" * 150)
                doc.add_paragraph("")  # Add some spacing
                iter += 1
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)  # Reset pointer to the start
            return output
    except Exception as e:
        pass
        #st.error(f"Error in function --> **get_docx_file**: {str(e)}")
        #st.warning(f"The questions below are not for the **{selected_file_type}** category. Please click on **Process Files** button to generate answers for the selected file type!")