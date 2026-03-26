import base64

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document

from utils.answer_question import (
    answer_question,
    answer_question_contract,
    answer_question_rfi_rfp,
    answer_question_shipping_logistics,
)
from utils.docx_file_ans import get_docx_file
from utils.embed_pdf import embed_file
from utils.embedding import load_embeddings
from utils.get_pdf_file import get_pdf_file
from utils.load_model import load_model
from utils.questions import (
    bid_rfi_rfp_questions,
    contract_questions,
    general_questions,
    shipping_logistics_questions,
)
from utils.text_file_ans import get_text_file


st.set_page_config(
    page_title="File Question Answering Tool",
    page_icon=":page_with_curl:",
    layout="wide",
)


def initialize_session_state():
    defaults = {
        "uploaded_files": [],
        "all_chunks": [],
        "all_file_page_mappings": [],
        "files_processed_check": False,
        "file_type_check": None,
        "answers_list": [],
        "files_pages_list": [],
        "answers_list_side": [],
        "files_pages_list_side": [],
        "ans_generated_check": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_index():
    import faiss

    st.session_state.index = faiss.IndexFlatL2(st.session_state.dimension)


def get_questions_by_type(selected_file_type):
    if selected_file_type == "General":
        return general_questions, answer_question
    elif selected_file_type == "Shipping & Logistics":
        return shipping_logistics_questions, answer_question_shipping_logistics
    elif selected_file_type == "RFI/RFP (Bid)":
        return bid_rfi_rfp_questions, answer_question_rfi_rfp
    elif selected_file_type == "Contract":
        return contract_questions, answer_question_contract
    else:
        return general_questions, answer_question


def display_pdf(selected_file):
    selected_file.seek(0)
    pdf_bytes = selected_file.getvalue()
    pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")

    html_code = f"""
        <div id="pdf-container" style="height:900px; overflow:auto;"></div>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
        <script>
        const pdfData = atob("{pdf_base64}");
        const pdfjsLib = window['pdfjs-dist/build/pdf'];
        pdfjsLib.GlobalWorkerOptions.workerSrc =
            'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';

        const loadingTask = pdfjsLib.getDocument({{data: pdfData}});
        loadingTask.promise.then(pdf => {{
            const container = document.getElementById('pdf-container');
            for (let i = 1; i <= pdf.numPages; i++) {{
                pdf.getPage(i).then(page => {{
                    const viewport = page.getViewport({{scale: 1.2}});
                    const canvas = document.createElement('canvas');
                    const context = canvas.getContext('2d');
                    canvas.height = viewport.height;
                    canvas.width = viewport.width;
                    container.appendChild(canvas);
                    page.render({{canvasContext: context, viewport: viewport}});
                }});
            }}
        }});
        </script>
    """
    components.html(html_code, height=920, scrolling=True)


def display_docx(selected_file):
    selected_file.seek(0)
    doc = Document(selected_file)
    st.markdown("### Document Content")
    has_content = False
    for para in doc.paragraphs:
        if para.text.strip():
            st.write(para.text)
            has_content = True
    if not has_content:
        st.info("No readable paragraph content found in this Word file.")


def display_excel(selected_file):
    selected_file.seek(0)
    excel_data = pd.read_excel(selected_file, sheet_name=None)
    if not excel_data:
        st.info("No sheets found in this Excel file.")
        return

    for sheet_name, df in excel_data.items():
        st.markdown(f"### Sheet: {sheet_name}")
        st.dataframe(df, use_container_width=True)


def display_selected_file(selected_file):
    file_name = selected_file.name.lower()

    if file_name.endswith(".pdf"):
        display_pdf(selected_file)
    elif file_name.endswith(".docx"):
        display_docx(selected_file)
    elif file_name.endswith((".xlsx", ".xls")):
        display_excel(selected_file)
    else:
        st.warning("Preview is not available for this file type.")


def process_uploaded_files():
    if not st.session_state.uploaded_files:
        st.warning("Please upload at least one file.")
        return

    st.session_state.all_chunks = []
    st.session_state.all_file_page_mappings = []
    st.session_state.answers_list = []
    st.session_state.files_pages_list = []
    st.session_state.answers_list_side = []
    st.session_state.files_pages_list_side = []
    st.session_state.ans_generated_check = False
    st.session_state.file_type_check = None
    reset_index()

    processed_any = False

    for uploaded_file in st.session_state.uploaded_files:
        with st.spinner(f"Processing {uploaded_file.name}..."):
            try:
                uploaded_file.seek(0)
                chunks, file_page_mapping = embed_file(uploaded_file, uploaded_file.name)

                if chunks:
                    st.session_state.all_chunks.extend(chunks)
                    st.session_state.all_file_page_mappings.extend(file_page_mapping)
                    processed_any = True
                else:
                    st.warning(f"No text could be extracted from {uploaded_file.name}")
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {e}")

    if processed_any:
        st.session_state.files_processed_check = True
        st.success("All supported files have been processed and indexed.")
    else:
        st.session_state.files_processed_check = False
        st.warning("No usable content was extracted from the uploaded files.")


def generate_answers_for_selected_type(selected_file_type):
    if not st.session_state.files_processed_check:
        st.warning("Please click 'Process Files' before generating answers.")
        return

    questions, answer_fn = get_questions_by_type(selected_file_type)

    with st.spinner("Generating answers..."):
        try:
            answers_list, files_pages_list = answer_fn(
                questions,
                st.session_state.all_chunks,
                st.session_state.all_file_page_mappings,
            )

            st.session_state.answers_list = answers_list or []
            st.session_state.files_pages_list = files_pages_list or []
            st.session_state.file_type_check = selected_file_type
            st.session_state.ans_generated_check = True

            if st.session_state.answers_list:
                st.success("Answers generated successfully.")
            else:
                st.warning("No answers were generated.")
        except Exception as e:
            st.error(f"Error generating answers: {e}")


def render_generated_answers(selected_file_type):
    if not st.session_state.ans_generated_check:
        return

    questions, _ = get_questions_by_type(selected_file_type)

    if not st.session_state.answers_list:
        st.warning("No answers available to display.")
        return

    st.markdown("---")
    st.subheader(f"Generated Answers - {st.session_state.file_type_check}")

    total_items = min(
        len(questions),
        len(st.session_state.answers_list),
        len(st.session_state.files_pages_list),
    )

    for i in range(total_items):
        st.markdown(f"### Question {i + 1}")
        st.write(questions[i])

        st.markdown("**Answer:**")
        st.write(st.session_state.answers_list[i])

        st.markdown("**Reference(s):**")
        refs = st.session_state.files_pages_list[i]
        if refs:
            for file_name, page_num in refs:
                st.write(f"- File: {file_name}, Page: {page_num}")
        else:
            st.write("No references available.")
        st.markdown("---")

    st.markdown("### Download Report")
    download_col1, download_col2, download_col3 = st.columns(3)

    try:
        txt_file = get_text_file(
            questions,
            st.session_state.answers_list,
            st.session_state.files_pages_list,
        )
        with download_col1:
            st.download_button(
                label="Download TXT",
                data=txt_file,
                file_name="question_answer_report.txt",
                mime="text/plain",
                use_container_width=True,
            )
    except Exception as e:
        st.error(f"Error creating TXT file: {e}")

    try:
        docx_file = get_docx_file(
            questions,
            st.session_state.answers_list,
            st.session_state.files_pages_list,
            st.session_state.file_type_check,
            selected_file_type,
        )
        with download_col2:
            st.download_button(
                label="Download DOCX",
                data=docx_file,
                file_name="question_answer_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    except Exception as e:
        st.error(f"Error creating DOCX file: {e}")

    try:
        pdf_file = get_pdf_file(
            questions,
            st.session_state.answers_list,
            st.session_state.files_pages_list,
            st.session_state.file_type_check,
            selected_file_type,
        )
        with download_col3:
            st.download_button(
                label="Download PDF",
                data=pdf_file,
                file_name="question_answer_report.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
    except Exception as e:
        st.error(f"Error creating PDF file: {e}")


def render_single_question_panel():
    st.subheader("Ask a Custom Question")

    if not st.session_state.files_processed_check:
        st.info("Upload and process files to ask a custom question.")
        return

    text_question = st.text_area("Enter your question", value="")

    if st.button("Generate Answer"):
        if not text_question.strip():
            st.warning("Please enter a question.")
        else:
            with st.spinner("Generating answer..."):
                try:
                    answers_list_side, files_pages_list_side = answer_question(
                        [text_question],
                        st.session_state.all_chunks,
                        st.session_state.all_file_page_mappings,
                    )
                    st.session_state.answers_list_side = answers_list_side or []
                    st.session_state.files_pages_list_side = files_pages_list_side or []
                except Exception as e:
                    st.error(f"Error generating custom answer: {e}")

    if st.session_state.answers_list_side:
        st.markdown("---")
        st.markdown("**Answer:**")
        st.write(st.session_state.answers_list_side[0])

        st.markdown("**Reference(s):**")
        refs = st.session_state.files_pages_list_side[0]
        if refs:
            for file_name, page_num in refs:
                st.write(f"- File: {file_name}, Page: {page_num}")
        else:
            st.write("No references available.")


def render_file_viewer():
    st.subheader("File Viewer")

    if not st.session_state.uploaded_files:
        st.info("Upload files to preview them here.")
        return

    file_names = ["Select a File"] + [file.name for file in st.session_state.uploaded_files]
    selected_file_name = st.selectbox("Select a File to view", file_names)

    if selected_file_name == "Select a File":
        return

    selected_file = None
    for uploaded_file in st.session_state.uploaded_files:
        if uploaded_file.name == selected_file_name:
            selected_file = uploaded_file
            break

    if selected_file is None:
        st.warning("Could not find the selected file.")
        return

    st.markdown(f"### Viewing File: {selected_file_name}")
    try:
        display_selected_file(selected_file)
    except Exception as e:
        st.error(f"Error displaying file: {e}")


def main_app():
    initialize_session_state()
    load_model()
    load_embeddings()

    st.markdown("<h1 style='text-align: center;'>File Question Answering Tool</h1>", unsafe_allow_html=True)
    st.markdown(
        "<p style='text-align: center;'>Upload PDF, Word, XLSX, or XLS files, process them, and generate answers.</p>",
        unsafe_allow_html=True,
    )

    main_col1, main_col2, main_col3 = st.columns([0.22, 0.40, 0.38])

    with main_col2:
        st.session_state.uploaded_files = st.file_uploader(
            "Upload Files",
            type=["pdf", "docx", "xlsx", "xls"],
            accept_multiple_files=True,
        )

        st.info("Click **Process Files** after uploading new files.")

        process_col1, process_col2 = st.columns(2)

        with process_col1:
            if st.button("Process Files", use_container_width=True):
                process_uploaded_files()

        file_types = ["General", "Shipping & Logistics", "RFI/RFP (Bid)", "Contract"]
        selected_file_type = process_col2.selectbox(
            "Select File Type",
            file_types,
        )

        if st.button("Generate Answers for Selected File Type", use_container_width=True):
            generate_answers_for_selected_type(selected_file_type)

        render_generated_answers(selected_file_type)

    with main_col1:
        render_single_question_panel()

    with main_col3:
        render_file_viewer()

    css = """
    <style>
        section.main > div {
            padding-bottom: 1rem;
        }
        [data-testid="stColumn"]:nth-child(2) > div > div {
            overflow-y: auto;
            max-height: 90vh;
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)


if __name__ == "__main__":
    main_app()
